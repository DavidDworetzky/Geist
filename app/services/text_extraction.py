import io
import re
import csv
import json
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Any
from docx import Document
from openpyxl import load_workbook
from PyPDF2 import PdfReader
import pdfplumber
import chardet

class TextExtractionService:
    """Service for extracting text content from various file types"""
    
    @staticmethod
    def extract_text_from_file(file_data: bytes, mime_type: str, filename: str) -> Dict[str, Any]:
        """
        Extract text from file based on MIME type and filename
        Returns dict with 'success' bool, 'text' content, and 'error' if any
        """
        try:
            # Get file extension
            file_ext = None
            if '.' in filename:
                file_ext = '.' + filename.rsplit('.', 1)[1].lower()
            
            # Route to appropriate extraction method
            if mime_type.startswith('text/') or file_ext in ['.txt', '.md']:
                return TextExtractionService._extract_from_text(file_data)
            elif mime_type == 'application/pdf' or file_ext == '.pdf':
                return TextExtractionService._extract_from_pdf(file_data)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_ext == '.docx':
                return TextExtractionService._extract_from_docx(file_data)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel'] or file_ext in ['.xlsx', '.xls']:
                return TextExtractionService._extract_from_excel(file_data)
            elif mime_type == 'text/csv' or file_ext == '.csv':
                return TextExtractionService._extract_from_csv(file_data)
            elif mime_type == 'application/json' or file_ext == '.json':
                return TextExtractionService._extract_from_json(file_data)
            elif mime_type in ['application/xml', 'text/xml'] or file_ext == '.xml':
                return TextExtractionService._extract_from_xml(file_data)
            else:
                return {
                    'success': False,
                    'error': f'Text extraction not supported for file type: {mime_type}'
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': f'Text extraction failed: {str(e)}'
            }
    
    @staticmethod
    def _extract_from_text(file_data: bytes) -> Dict[str, Any]:
        """Extract text from plain text files with encoding detection"""
        # Detect encoding
        detected = chardet.detect(file_data)
        encoding = detected.get('encoding', 'utf-8')
        
        try:
            text = file_data.decode(encoding)
            return {'success': True, 'text': text}
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            try:
                text = file_data.decode('utf-8', errors='replace')
                return {'success': True, 'text': text}
            except Exception as e:
                return {'success': False, 'error': f'Failed to decode text: {str(e)}'}
    
    @staticmethod
    def _extract_from_pdf(file_data: bytes) -> Dict[str, Any]:
        """Extract text from PDF files using multiple methods"""
        text_content = []
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            if text_content:
                return {'success': True, 'text': '\n\n'.join(text_content)}
            
            # Method 2: Fallback to PyPDF2
            pdf_reader = PdfReader(io.BytesIO(file_data))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            
            if text_content:
                return {'success': True, 'text': '\n\n'.join(text_content)}
            else:
                return {'success': False, 'error': 'No text could be extracted from PDF'}
                
        except Exception as e:
            return {'success': False, 'error': f'PDF text extraction failed: {str(e)}'}
    
    @staticmethod
    def _extract_from_docx(file_data: bytes) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            doc = Document(io.BytesIO(file_data))
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            text = '\n\n'.join(paragraphs)
            return {'success': True, 'text': text}
            
        except Exception as e:
            return {'success': False, 'error': f'DOCX text extraction failed: {str(e)}'}
    
    @staticmethod
    def _extract_from_excel(file_data: bytes) -> Dict[str, Any]:
        """Extract text from Excel files"""
        try:
            workbook = load_workbook(io.BytesIO(file_data), data_only=True)
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = [f"Sheet: {sheet_name}"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip() and row_text != ' |  |  |  |  |  |  |  |  |  |':
                        sheet_content.append(row_text)
                
                if len(sheet_content) > 1:  # More than just the sheet name
                    text_content.extend(sheet_content)
                    text_content.append('')  # Empty line between sheets
            
            text = '\n'.join(text_content)
            return {'success': True, 'text': text}
            
        except Exception as e:
            return {'success': False, 'error': f'Excel text extraction failed: {str(e)}'}
    
    @staticmethod
    def _extract_from_csv(file_data: bytes) -> Dict[str, Any]:
        """Extract text from CSV files"""
        try:
            # Detect encoding
            detected = chardet.detect(file_data)
            encoding = detected.get('encoding', 'utf-8')
            
            text_data = file_data.decode(encoding)
            csv_reader = csv.reader(io.StringIO(text_data))
            
            rows = []
            for row in csv_reader:
                row_text = ' | '.join(row)
                rows.append(row_text)
            
            text = '\n'.join(rows)
            return {'success': True, 'text': text}
            
        except Exception as e:
            return {'success': False, 'error': f'CSV text extraction failed: {str(e)}'}
    
    @staticmethod
    def _extract_from_json(file_data: bytes) -> Dict[str, Any]:
        """Extract text from JSON files"""
        try:
            # Detect encoding
            detected = chardet.detect(file_data)
            encoding = detected.get('encoding', 'utf-8')
            
            text_data = file_data.decode(encoding)
            json_data = json.loads(text_data)
            
            # Convert JSON to readable text format
            text = json.dumps(json_data, indent=2, ensure_ascii=False)
            return {'success': True, 'text': text}
            
        except Exception as e:
            return {'success': False, 'error': f'JSON text extraction failed: {str(e)}'}
    
    @staticmethod
    def _extract_from_xml(file_data: bytes) -> Dict[str, Any]:
        """Extract text from XML files"""
        try:
            # Detect encoding
            detected = chardet.detect(file_data)
            encoding = detected.get('encoding', 'utf-8')
            
            text_data = file_data.decode(encoding)
            
            # Parse XML and extract text content
            root = ET.fromstring(text_data)
            text_content = []
            
            def extract_text_recursive(element):
                if element.text and element.text.strip():
                    text_content.append(element.text.strip())
                for child in element:
                    extract_text_recursive(child)
                if element.tail and element.tail.strip():
                    text_content.append(element.tail.strip())
            
            extract_text_recursive(root)
            
            # Also include formatted XML for structure
            formatted_xml = ET.tostring(root, encoding='unicode')
            text = f"Extracted Text:\n{' '.join(text_content)}\n\nXML Structure:\n{formatted_xml}"
            
            return {'success': True, 'text': text}
            
        except Exception as e:
            return {'success': False, 'error': f'XML text extraction failed: {str(e)}'}
    
    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        # Limit text length to prevent database issues (e.g., 50KB)
        max_length = 50000
        if len(text) > max_length:
            text = text[:max_length] + "\n\n[Text truncated due to length]"
        
        return text